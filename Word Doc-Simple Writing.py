import docx
import os
from docx.shared import Inches
doc = docx.Document()
doc.add_heading('Test Doc',0)
parag=doc.add_paragraph("Hello!")
parag.add_run("This word document was created using, ")
parag.add_run("Python").bold=True
doc.add_heading('Heading Level 1',1)
doc.add_heading('Heading Level 2',2)
italicparagraph=doc.add_paragraph()
italicparagraph.add_run("This line is in italics!").italic=True
doc.add_picture('sample_image.jpg',width=Inches(2.5))
doc.save("test.docx")
os.system("start test.docx")
