import docx
import os

doc = docx.Document()

doc.add_heading('Table Document',0)

records = [
	[1, 'Chicken Pasta Bake', '£4.50'],
	[2, 'Starfish Cakes', '£5.00'],
	[3, 'Cod Confit', '£2.75']
		   ]
menuTable = doc.add_table(rows=1,cols=3)
menuTable.style= 'Medium Shading 2 Accent 3'
hdr_Cells = menuTable.rows[0].cells
hdr_Cells[0].text = 'ID'
hdr_Cells[1].text = 'Meal Name'
hdr_Cells[2].text = 'Price'

for ID, nameOfMeal, price in records:
	row_Cells = menuTable.add_row().cells
	row_Cells[0].text= str(ID)
	row_Cells[1].text = nameOfMeal
	row_Cells[2].text = price
	
doc.save('table.docx')
os.system("start table.docx")
